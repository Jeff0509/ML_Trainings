import requests
from bs4 import BeautifulSoup
from docx import Document
from PyPDF2 import PdfReader
import io

def fetch_html(url):
    try:
        response = requests.get(url, timeout=10)
        response.raise_for_status()
        return response.text
    except requests.exceptions.RequestException as e:
        print(f"An error occurred while fetching HTML: {e}")
        return None

def html_to_text(html):
    soup = BeautifulSoup(html, 'html.parser')
    # Remove script and style elements
    for script_or_style in soup(['script', 'style']):
        script_or_style.decompose()
    # Get text
    text = soup.get_text()
    # Break into lines and remove leading/trailing whitespace
    lines = (line.strip() for line in text.splitlines())
    # Break multi-headlines into a line each
    chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
    # Drop blank lines
    text = '\n'.join(chunk for chunk in chunks if chunk)
    return text

def fetch_content(url):
    try:
        response = requests.get(url, timeout=10)
        response.raise_for_status()
        if response.status_code == 200:
            return response.content, response.headers.get('content-type')
        else:
            return None, None
    except requests.exceptions.RequestException as e:
        print(f"An error occurred while fetching content: {e}")
        return None, None

def extract_text_from_pdf(pdf_content):
    pdf_text = ""
    pdf_reader = PdfReader(io.BytesIO(pdf_content))
    for page in pdf_reader.pages:
        pdf_text += page.extract_text()
    return pdf_text

def create_word_document(urls, filename):
    doc = Document()
    for url in urls:
        content, content_type = fetch_content(url)
        if content and content_type:
            if 'text/html' in content_type:
                # If the content is HTML, convert it to text
                text = html_to_text(content.decode('utf-8'))
                doc.add_heading(url, level=1)
                doc.add_paragraph(text)
            elif 'application/pdf' in content_type:
                # If the content is a PDF document, extract text from it
                pdf_text = extract_text_from_pdf(content)
                doc.add_heading(url, level=1)
                doc.add_paragraph(pdf_text)
            else:
                # If the content is not HTML or PDF, save it as-is
                with open('temp_file', 'wb') as file:
                    file.write(content)
                # Append the non-HTML/PDF content to the Word document
                doc2 = Document('temp_file')
                for element in doc2.element.body:
                    doc.element.body.append(element)
        doc.add_page_break()

    doc.save(filename)

def search_google_news(query, api_key, cse_id, from_date):
    search_url = "https://www.googleapis.com/customsearch/v1"
    params = {
        'q': query,
        'cx': cse_id,
        'key': api_key,
        'dateRestrict': f'd{from_date}',
        'gl': 'au',
        'num': 7,  # Increase the number of results per request
    }

    all_results = []  # List to store all results

    while True:
        response = requests.get(search_url, params=params)
        data = response.json()
        items = data.get('items', [])

        if not items:
            break  # No more results to fetch

        all_results.extend(items)

        # Check if there are more results to fetch
        if 'nextPage' in data.get('queries', {}).get('nextPage', [{}])[0]:
            params['start'] = data['queries']['nextPage'][0]['startIndex']
        else:
            break

    return all_results

def main():
    api_key = 'AIzaSyBX-V1qjCjwGhykzxaiYkyO2HK0RjL7Fr8'
    cse_id = '237845578820a4aa8'
    #query = 'Vocational Education and Training Victoria Australia'
    query = 'Xi Jinping'
    filename = query + '.docx'
    from_date = '20240101'

    results = search_google_news(query, api_key, cse_id, from_date)
    all_html = ""
    
    urls = []  # Create an empty list to store the URLs

    for item in results.get('items', []):
        url = item['link']
        urls.append(url)
        #html_content = fetch_html(url)
        #all_html += f"\n<!-- Start of content from {url} -->\n"
        #all_html += html_content
        #all_html += f"\n<!-- End of content from {url} -->\n"

    create_word_document(urls,filename)
    #with open("combined_news.html", "w", encoding="utf-8") as file:
    #    file.write(all_html)

    print(f"Combined document created as '{filename}'")

if __name__ == "__main__":
    main()